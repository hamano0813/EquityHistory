# !/usr/bin/env python
# -*- coding: utf-8 -*-

# noinspection PyPackageRequirements
import docx
from lxml import etree

TITLE = ('序号', '股东名称', '注册资本额\n（万元）', '持股比例\n（%）')


# noinspection PyPep8Naming
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = docx.oxml.OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            element = tcBorders.find(docx.oxml.ns.qn(tag))
            if element is None:
                element = docx.oxml.OxmlElement(tag)
                tcBorders.append(element)

            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(docx.oxml.ns.qn('w:{}'.format(key)), str(edge_data[key]))


def set_table_border(table):
    rows = len(table.rows)
    cols = len(table.columns)

    for rid, row in enumerate(table.rows):
        for cid, cell in enumerate(row.cells):
            setting = {
                'start': {'sz': 6, 'val': 'dotted'},
                'top': {'sz': 6, 'val': 'dotted'},
                'end': {'sz': 6, 'val': 'dotted'},
                'bottom': {'sz': 6, 'val': 'dotted'},
            }
            if rid == 0:
                setting.update({'top': {'sz': 12, 'val': 'single'}})
            elif rid == rows - 1:
                setting.update({'bottom': {'sz': 12, 'val': 'single'}})
            if cid == 0:
                setting.update({'start': {'sz': 12, 'val': 'nil'}})
            elif cid == cols - 1:
                setting.update({'end': {'sz': 12, 'val': 'nil'}})
            set_cell_border(cell, **setting)


def convert_docx(groups, company, file_path):
    doc = docx.Document()
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'宋体')

    doc.add_paragraph().add_run('股权结构及历史沿革').bold = True

    for group in groups.values():
        overview = group['overview']
        shareholding = group['shareholding']

        overview_lines = overview.split('\n')
        doc.add_paragraph().add_run(overview_lines[0]).bold = True
        doc.add_paragraph().add_run(overview_lines[1])

        table = doc.add_table(rows=len(shareholding) + 1, cols=4)
        set_table_border(table)
        table.cell(len(shareholding), 0).merge(table.cell(len(shareholding), 1))
        table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER

        widths = (docx.shared.Cm(1.25), docx.shared.Cm(10), docx.shared.Cm(2.75), docx.shared.Cm(2))
        for rid, row in enumerate(table.rows):
            row.height = docx.shared.Cm(0.7)
            for cid, cell in enumerate(row.cells):
                paragraph = cell.paragraphs[0]
                paragraph_format = paragraph.paragraph_format
                paragraph_format.line_spacing = 1.0
                paragraph_format.space_after = docx.shared.Pt(0)
                if rid == 0:
                    run = paragraph.add_run(TITLE[cid])
                    run.bold = True
                    paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                else:
                    run = paragraph.add_run(shareholding[rid - 1][cid])
                    if cid in (2, 3):
                        paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                    elif not cid:
                        paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    elif rid == len(shareholding):
                        paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    if rid == len(shareholding):
                        run.bold = True
                run.font.size = docx.shared.Pt(10.5)
                cell.vertical_alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                cell.width = widths[cid]

    doc.add_paragraph().add_run(f'截至评估基准日，{company}股权结构与上表一致，无变化。')

    for paragraph in doc.paragraphs:
        if paragraph._element.xpath('.//w:tbl'):
            continue
        for run in paragraph.runs:
            run.font.size = docx.shared.Pt(12)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format.line_spacing = 1.5
        paragraph_format.first_line_indent = 266700
        paragraph_format.space_after = docx.shared.Pt(0)

    tblw = doc._element.xpath('//w:tblW')
    for tblw_o in tblw:
        tblw_n = etree.fromstring(etree.tostring(tblw_o))
        tblw_n.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'pct')
        tblw_n.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', '5000')
        tblw_o.getparent().replace(tblw_o, tblw_n)

    doc.save(file_path)
